from docx import Document
from docx.shared import Cm, Pt

from docx.document import Document as Doc

# 创建代表Word文档的Doc对象
document = Document()

# 添加大标题
document.add_heading('快快乐乐学Python', 0)
# 添加段落  add_paragraph('文字') 本身就会创建一个带文字的 run。
# paragraph 负责“换不换段”，run 负责“这一小段文字长什么样”。
p = document.add_paragraph('1、Python是一门非常流行的编程语言，它{name}')
# 等价于
p = document.add_paragraph()
p.add_run('2、Python是一门非常流行的编程语言，它{name}')
run = p.add_run('简单')
run.bold = True
run.font.size = Pt(18)
p.add_run('而且')
p = document.add_paragraph('3、Python是一门非常流行的编程语言，2')
run = p.add_run('优雅')
run.bold = True
run.font.size = Pt(18)
p.add_run('。')
p.add_run('我的名字是')
run = p.add_run('{name}')
run.bold = True
run.font.size = Pt(18)



# 添加一级标题
document.add_heading('Heading, level 1', level=1)
# 添加带样式的段落
document.add_paragraph('Intense quote', style='Intense Quote')
# 添加无序列表
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'second item in unordered list', style='List Bullet'
)

# 添加有序列表
document.add_paragraph(
    'first item in unordered list', style='List Number'
)
document.add_paragraph(
    'second item in unordered list', style='List Number'
)

# 添加分节符
document.add_section()

records = (
    ('骆昊', '男', '1995-5-5'),
    ('孙美丽', '女', '1992-2-2')    
)

# 添加表格
table = document.add_table(rows=1, cols=3)
table.style = 'Dark List'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '姓名'
hdr_cells[1].text = '性别'
hdr_cells[2].text = '出生日期'
# 为表格添加行
for name, sex, birthday in records:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = sex
    row_cells[2].text = birthday

# 添加分页符
document.add_page_break()

# 保存文档
document.save('demo.docx')



doc = Document('demo.docx')
for no, p in enumerate(doc.paragraphs):
    print(no, p.text)

for p in doc.paragraphs:
    for run in p.runs:
        run.text = run.text.replace('{name}', '张三')

for no, p in enumerate(doc.paragraphs):
    print(no, p.text)